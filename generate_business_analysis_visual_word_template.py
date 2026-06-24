#!/usr/bin/env python3
"""Generate a deep-blue visual Word template for business analysis reports."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "business-analysis-visual-word-template.docx"

NAVY = "071A2F"
NAVY_2 = "102B46"
BLUE = "1E66D0"
BLUE_SOFT = "DCEBFF"
BLUE_TINT = "EAF2FF"
BLUE_MIST = "F3F7FC"
INK = "262626"
GRAY = "617082"
WHITE = "FFFFFF"
ICE = "7FB6FF"


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_font(run, name: str = "Aptos", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.font.bold = bold


def cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color: str = "D9DEE4", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def row_height(row, inches: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    height = tr_pr.find(qn("w:trHeight"))
    if height is None:
        height = OxmlElement("w:trHeight")
        tr_pr.append(height)
    height.set(qn("w:val"), str(int(inches * 1440)))
    height.set(qn("w:hRule"), "atLeast")


def clear_cell(cell, fill: str = WHITE, border: str = "D9DEE4") -> None:
    cell_shading(cell, fill)
    cell_border(cell, border)
    cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def add_cell_text(cell, text: str, *, size: float = 8.5, color: str = INK, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, font: str = "Aptos") -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_font(run, font, size=size, color=color, bold=bold)


def add_doc_text(document: Document, text: str, *, size: float = 10, color: str = INK, bold: bool = False, font: str = "Aptos", align=WD_ALIGN_PARAGRAPH.LEFT, before: float = 0, after: float = 6) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_font(run, font, size=size, color=color, bold=bold)


def setup_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    h1 = document.styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = rgb(NAVY)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(6)

    h2 = document.styles["Heading 2"]
    h2.font.name = "Aptos Display"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLUE)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("[Client / Company]  |  Business Analysis")
    set_font(run, size=8.2, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Facts  |  Findings  |  Implications  |  Actions")
    set_font(run, size=8, color=GRAY)

    document.core_properties.title = "Business Analysis Visual Word Template"
    document.core_properties.subject = "Business analysis report template aligned to the business-analysis skill"
    document.core_properties.author = "OpenAI Codex"
    return document


def add_colored_bar(document: Document, label: str, value: str, fill: str = NAVY) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.05)
    table.columns[1].width = Inches(6.0)
    label_cell, value_cell = table.rows[0].cells
    clear_cell(label_cell, BLUE, BLUE)
    clear_cell(value_cell, fill, fill)
    add_cell_text(label_cell, label, size=16, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="Aptos Display")
    add_cell_text(value_cell, value, size=13, color=WHITE, bold=True, font="Aptos Display")


def add_note_box(document: Document, title: str, body: str, fill: str = BLUE_SOFT, stripe: str = BLUE) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.16)
    table.columns[1].width = Inches(6.9)
    stripe_cell, body_cell = table.rows[0].cells
    clear_cell(stripe_cell, stripe, stripe)
    cell_margins(stripe_cell, top=30, bottom=30, start=20, end=20)
    clear_cell(body_cell, fill, "C9D7EA")
    cell_margins(body_cell, top=120, bottom=120, start=140, end=140)
    add_cell_text(body_cell, title, size=11, color=NAVY, bold=True, font="Aptos Display")
    add_cell_text(body_cell, body, size=9.4, color=INK)


def add_cards(document: Document, cards: list[tuple[str, str, str]], cols: int = 3) -> None:
    rows = (len(cards) + cols - 1) // cols
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            clear_cell(cell, BLUE_MIST, "DDE3EA")
            cell_margins(cell, top=130, bottom=130, start=130, end=130)
    for index, (label, value, body) in enumerate(cards):
        cell = table.cell(index // cols, index % cols)
        clear_cell(cell, BLUE_TINT if index % 2 == 0 else BLUE_MIST, "DDE3EA")
        add_cell_text(cell, label.upper(), size=7.8, color=BLUE, bold=True)
        add_cell_text(cell, value, size=13, color=NAVY, bold=True, font="Aptos Display")
        if body:
            add_cell_text(cell, body, size=8.4, color=INK)


def add_template_table(document: Document, headers: list[str], rows: list[list[str]], header_fill: str = NAVY) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        clear_cell(table.rows[0].cells[index], header_fill, header_fill)
        add_cell_text(table.rows[0].cells[index], header, size=8.2, color=WHITE, bold=True)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 == 0 else BLUE_MIST
        for index, value in enumerate(row_data):
            clear_cell(cells[index], fill, "D9DEE4")
            add_cell_text(cells[index], value, size=8.2, color=INK)


def add_flow_diagram(document: Document, steps: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = 7.0 / max(len(steps), 1)
    for index, (label, body) in enumerate(steps):
        cell = table.rows[0].cells[index]
        table.columns[index].width = Inches(width)
        fill = NAVY if index == 0 else NAVY_2 if index % 2 else BLUE_SOFT
        text_color = WHITE if fill in [NAVY, NAVY_2] else NAVY
        clear_cell(cell, fill, BLUE)
        cell_margins(cell, top=130, bottom=120, start=100, end=100)
        add_cell_text(cell, label, size=10.2, color=ICE if text_color == WHITE else BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="Aptos Display")
        add_cell_text(cell, body, size=8, color=text_color, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_narrative_block(document: Document, claim: str, prompts: list[str]) -> None:
    document.add_heading("Narrative Draft Space", level=2)
    items = [
        ("Opening finding", f"[Write 4-6 polished sentences that make the section's main finding: {claim}]"),
        ("Evidence base", f"[State verified facts, source status, and assumptions. Address: {prompts[0] if prompts else 'what must be proven'}]"),
        ("Analysis logic", "[Explain the business logic, tradeoff, root cause, financial driver, process constraint, or stakeholder implication.]"),
        ("Implication", f"[Explain what this means for the decision, roadmap, risk, or next validation. Address: {prompts[1] if len(prompts) > 1 else 'the decision implication'}]"),
    ]
    for label, body in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.12
        run = paragraph.add_run(f"{label}: ")
        set_font(run, "Aptos Display", size=10.2, color=NAVY, bold=True)
        run = paragraph.add_run(body)
        set_font(run, "Aptos", size=9.8, color=INK)


def add_section_visual(document: Document, visual: str | None) -> None:
    if visual == "summary":
        document.add_heading("Executive Decision Snapshot", level=2)
        add_cards(document, [("Key finding", "[Finding]", "What matters most."), ("Decision", "[Action]", "What to do."), ("Risk", "[Risk]", "What could fail.")])
    elif visual == "method":
        document.add_heading("Analysis Flow", level=2)
        add_flow_diagram(document, [("Question", "[Decision]"), ("Evidence", "[Sources]"), ("Analysis", "[Method]"), ("Action", "[Recommendation]")])
    elif visual == "market":
        document.add_heading("Context Flow", level=2)
        add_flow_diagram(document, [("External", "[Driver]"), ("Customer", "[Need]"), ("Business", "[Constraint]"), ("Impact", "[Implication]")])
    elif visual == "operating":
        document.add_heading("Operating Model Flow", level=2)
        add_flow_diagram(document, [("Input", "[Resource]"), ("Process", "[Activity]"), ("Output", "[Deliverable]"), ("Metric", "[KPI]")])
    elif visual == "roadmap":
        document.add_heading("Implementation Roadmap", level=2)
        add_flow_diagram(document, [("0-30D", "[Stabilize]"), ("31-90D", "[Improve]"), ("3-6M", "[Scale]"), ("6M+", "[Embed]")])
    elif visual == "quality":
        document.add_heading("Quality Gate", level=2)
        add_template_table(document, ["Area", "OK", "Issue", "Fix"], [["Facts", "[ ]", "[ ]", "[ ]"], ["Sources", "[ ]", "[ ]", "[ ]"], ["Metrics", "[ ]", "[ ]", "[ ]"], ["Actions", "[ ]", "[ ]", "[ ]"]])


def add_section_page(
    document: Document,
    number: str,
    title: str,
    claim: str,
    prompts: list[str],
    headers: list[str],
    rows: list[list[str]],
    visual: str | None = None,
) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    add_colored_bar(document, number, title)
    document.add_heading(f"{number} {title}", level=1)
    add_note_box(document, "Lead finding", claim, fill=BLUE_TINT, stripe=BLUE)
    add_narrative_block(document, claim, prompts)
    add_section_visual(document, visual)
    add_template_table(document, headers, rows)


def add_cover_page(document: Document) -> None:
    add_colored_bar(document, "BA", "Business Analysis Report")
    add_doc_text(document, "BUSINESS ANALYSIS", size=9.2, color=BLUE, bold=True, before=36, after=4)
    add_doc_text(document, "Business Analysis\nReport", size=54, color=NAVY, bold=True, font="Aptos Display", after=12)
    add_doc_text(
        document,
        "[Template for market, operational, financial, process, capability, and strategy analysis reports.]",
        size=14.5,
        color=INK,
        bold=True,
        after=34,
    )
    add_cards(
        document,
        [
            ("Business question", "[What must be answered?]", "Define the decision first."),
            ("Audience", "[Board / leadership / operations]", "Adapt depth and tone."),
            ("Evidence standard", "[Sources / data / assumptions]", "Mark unknowns explicitly."),
        ],
    )
    add_doc_text(document, "", after=34)
    add_note_box(
        document,
        "Core rule",
        "This is not a generic summary. Every finding should connect facts to implications, recommendations, owners, metrics, and next validation.",
        fill=BLUE_SOFT,
        stripe=BLUE,
    )


def generate(output_path: Path = OUTPUT_PATH) -> Path:
    document = setup_document()
    add_cover_page(document)

    sections = [
        ("01", "Executive Summary", "Summarize the business question, key finding, recommendation, and decision implication.", ["What is the answer in one page?", "What should the audience do next?"], ["Item", "Executive Content"], [["Business question", "[Question]"], ["Key finding", "[Finding]"], ["Recommendation", "[Action]"], ["Confidence / caveat", "[Confidence and limitation]"]], "summary"),
        ("02", "Decision Goal And Audience", "The analysis must define who will use it and what decision they need to make.", ["Who is the audience?", "What action should they take after reading?"], ["Field", "Working Definition"], [["Decision goal", "[Decision / approval / diagnosis / prioritization]"], ["Audience", "[Board / leadership / investor / operations / product]"], ["Scope boundary", "[Included / excluded]"], ["Time period", "[Period / date cut-off]"]], None),
        ("03", "Business Context And Scope", "Context establishes the company, unit, geography, process, market, and constraints under analysis.", ["What is the business setting?", "What is inside and outside scope?"], ["Context Area", "Definition", "Implication"], [["Company / unit", "[Name / business unit]", "[Why relevant]"], ["Geography", "[Market / region]", "[Localization implication]"], ["Business problem", "[Problem statement]", "[Decision implication]"], ["Constraint", "[Budget / time / regulation]", "[Risk]"]], "market"),
        ("04", "Methodology And Source Register", "The report should show how the analysis was conducted and where the evidence came from.", ["Which method was used?", "Which sources are verified or missing?"], ["Method / Source", "Use", "Output / Treatment"], [["Problem statement", "[Define need and boundary]", "[Scope clarity]"], ["Process map / SIPOC", "[Operational diagnosis]", "[Bottleneck view]"], ["KPI tree / ROI", "[Financial logic]", "[Metric impact]"], ["Source register", "[Evidence chain]", "[Confidence and gaps]"]], "method"),
        ("05", "Current-State Findings", "Current-state findings describe what is happening now, supported by facts and evidence.", ["What is the current state?", "Which facts support the finding?"], ["Finding", "Evidence", "Implication", "Status"], [["[Finding 1]", "[Source / data]", "[Implication]", "[Verified / TBD]"], ["[Finding 2]", "[Source / data]", "[Implication]", "[Verified / TBD]"], ["[Finding 3]", "[Source / data]", "[Implication]", "[Verified / TBD]"]], None),
        ("06", "Market / Customer / Industry Context", "Use this section when external market, customer, or industry conditions affect the decision.", ["Which external forces matter?", "How do they change the recommendation?"], ["External Factor", "Evidence", "Business Impact"], [["[Customer need]", "[Source]", "[Impact]"], ["[Market trend]", "[Source]", "[Impact]"], ["[Industry pressure]", "[Source]", "[Impact]"]], "market"),
        ("07", "Process / Operating Model / Capability Analysis", "Analyze how work gets done and where process, capability, or operating model gaps sit.", ["Where are bottlenecks?", "Which capabilities are missing or weak?"], ["Area", "Current State", "Gap", "Required Capability"], [["[Process step]", "[Current]", "[Gap]", "[Need]"], ["[Operating model]", "[Current]", "[Gap]", "[Need]"], ["[Capability]", "[Current]", "[Gap]", "[Need]"]], "operating"),
        ("08", "Financial / KPI / Scenario Analysis", "Define the numeric logic, assumptions, scenarios, and KPI implications.", ["Which metrics drive the decision?", "What assumptions change the outcome?"], ["Metric", "Definition", "Unit / Currency", "Formula / Assumption", "Source"], [["[KPI]", "[Definition]", "[Unit]", "[Formula]", "[Source]"], ["[Cost driver]", "[Definition]", "[Currency]", "[Formula]", "[Source]"], ["[Scenario]", "[Definition]", "[Period]", "[Assumption]", "[Source]"]], None),
        ("09", "Gap Analysis And Root Causes", "Translate findings into gaps and root causes rather than surface symptoms.", ["What is the real gap?", "What is the underlying cause?"], ["Gap", "Evidence", "Root Cause", "Business Impact"], [["[Gap 1]", "[Evidence]", "[Root cause]", "[Impact]"], ["[Gap 2]", "[Evidence]", "[Root cause]", "[Impact]"], ["[Gap 3]", "[Evidence]", "[Root cause]", "[Impact]"]], None),
        ("10", "Options Or Recommendations", "Compare feasible options and recommend a clear path forward.", ["Which options were considered?", "Why is the recommended option best?"], ["Option", "Benefit", "Tradeoff", "Recommendation"], [["[Option A]", "[Benefit]", "[Tradeoff]", "[Recommended / not]"], ["[Option B]", "[Benefit]", "[Tradeoff]", "[Recommended / not]"], ["[Option C]", "[Benefit]", "[Tradeoff]", "[Recommended / not]"]], None),
        ("11", "Implementation Roadmap", "Turn the recommendation into phases, owners, milestones, and success metrics.", ["What happens first?", "Who owns each action and how is success measured?"], ["Phase", "Action", "Owner / Function", "Metric", "Dependency"], [["0-30D", "[Action]", "[Owner]", "[Metric]", "[Dependency]"], ["31-90D", "[Action]", "[Owner]", "[Metric]", "[Dependency]"], ["3-6M", "[Action]", "[Owner]", "[Metric]", "[Dependency]"]], "roadmap"),
        ("12", "Risks, Assumptions, And Limitations", "Make uncertainty visible and define the next validation steps.", ["Which assumptions could change the recommendation?", "What needs validation?"], ["Risk / Assumption", "Why It Matters", "Current Treatment", "Next Validation"], [["[Risk]", "[Impact]", "[Mitigation / caveat]", "[Validation]"], ["[Assumption]", "[Impact]", "[Scenario]", "[Validation]"], ["[Limitation]", "[Impact]", "[TBD]", "[Validation]"]], None),
        ("13", "Appendices / Source List", "Provide the source list, variable register, metric dictionary, and supporting exhibits.", ["Where did each key claim come from?", "What source gaps remain?"], ["Claim / Data Point", "Source", "Type", "Date", "Confidence"], [["[Claim]", "[URL / file]", "[Official / user / third-party]", "[Date]", "[High / medium / low]"], ["[Metric]", "[Source]", "[Dashboard / spreadsheet]", "[Date]", "[Confidence]"], ["[Assumption]", "[Source]", "[User / inference]", "[Date]", "[Needs verification]"]], "quality"),
    ]

    for args in sections:
        add_section_page(document, *args)

    document.save(output_path)
    return output_path


def main() -> None:
    path = generate()
    print(f"Generated {path}")


if __name__ == "__main__":
    main()
